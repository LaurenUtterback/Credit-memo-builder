"""Turn uploaded files into Anthropic message content blocks.

Every uploader in this app (credit memo, participation agreement, loan
documents) used to hand the browser's reported MIME type straight to the API as
a document block's ``media_type``. The API accepts only ``application/pdf``
there, so anything else failed the whole extraction with a raw 400:

    messages.0.content.3.document.source.base64.media_type:
        Input should be 'application/pdf'

Two things went wrong in that design. Browsers report no useful MIME for many
files pulled off a network share (Windows hands over ``application/octet-stream``
or an empty string), so even a genuine PDF could be rejected. And a Word or
Excel file — routinely part of a deal package — had no path through at all.

So the file's own magic bytes decide what it is, never the declared MIME, and
each kind gets the block the API actually accepts:

    PDF                  -> document block
    PNG/JPEG/GIF/WEBP    -> image block
    DOCX / XLSX          -> text block (converted here with python-docx/openpyxl)
    txt/csv/md/json      -> text block

Anything else raises a RuntimeError naming the file, so the user is told which
document to convert instead of seeing an API error code.
"""

from __future__ import annotations

import base64
import csv
import io
import json
import os

from .models import UploadedDoc


# Cap converted spreadsheets so one oversized workbook can't crowd the deal
# documents out of the context window.
_MAX_SHEET_ROWS = 400

_IMAGE_MAGIC = (
    (b"\x89PNG\r\n\x1a\n", "image/png"),
    (b"\xff\xd8\xff", "image/jpeg"),
    (b"GIF87a", "image/gif"),
    (b"GIF89a", "image/gif"),
)

_TEXT_EXTS = {".txt", ".csv", ".md", ".json", ".tsv"}


def _head(doc: UploadedDoc, nbytes: int = 48) -> bytes:
    """Decode just the first few bytes of the payload to sniff its type."""
    chunk = (doc.b64 or "")[: (nbytes // 3 + 1) * 4]
    chunk = chunk[: len(chunk) // 4 * 4]          # base64 decodes in 4-char groups
    if not chunk:
        return b""
    try:
        return base64.b64decode(chunk)
    except Exception:                              # noqa: BLE001 - malformed upload
        return b""


def _raw(doc: UploadedDoc) -> bytes:
    try:
        return base64.b64decode(doc.b64 or "")
    except Exception as exc:                       # noqa: BLE001
        raise RuntimeError(f"{doc.filename}: the upload is not valid base64.") from exc


def _ext(doc: UploadedDoc) -> str:
    return os.path.splitext(doc.filename or "")[1].lower()


def _docx_text(data: bytes) -> str:
    """Flatten a .docx to text: paragraphs first, then every table row."""
    import docx                                    # provided by python-docx

    d = docx.Document(io.BytesIO(data))
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _xlsx_text(data: bytes) -> str:
    """Flatten a .xlsx to text, one ``a | b | c`` line per non-empty row.

    Values, not formulas (``data_only``) — a formula string tells Claude nothing
    about the number the underwriter sees in the sheet.
    """
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    out: list[str] = []
    try:
        for ws in wb.worksheets:
            out.append(f"--- sheet: {ws.title} ---")
            for n, row in enumerate(ws.iter_rows(values_only=True)):
                if n >= _MAX_SHEET_ROWS:
                    out.append(f"[... truncated at {_MAX_SHEET_ROWS} rows ...]")
                    break
                cells = ["" if v is None else str(v).strip() for v in row]
                if any(cells):
                    out.append(" | ".join(cells).rstrip(" |"))
    finally:
        wb.close()
    return "\n".join(out)


def _plain_text(data: bytes, ext: str) -> str:
    text = data.decode("utf-8", errors="replace")
    if ext == ".json":
        try:                                       # pretty-print so it reads well
            return json.dumps(json.loads(text), indent=1)[:200_000]
        except ValueError:
            return text
    if ext in {".csv", ".tsv"}:
        delim = "\t" if ext == ".tsv" else ","
        rows = list(csv.reader(io.StringIO(text), delimiter=delim))
        return "\n".join(" | ".join(r) for r in rows if any(c.strip() for c in r))
    return text


# --- oversized uploads: shrink instead of refuse ------------------------------
# The Messages API hard-caps a request at ~32 MB, so an over-budget upload used
# to be refused outright (naming the largest files). Lauren, 2026-08-14: shrink
# it to fit instead — largest blocks first, and only as much as needed:
#   pass 1: re-encode the images INSIDE a PDF (downscale + JPEG). A 300-DPI
#           scanned bundle shrinks several-fold and stays a visual document.
#   pass 2: swap a PDF for its text layer (like the .docx/.xlsx handling) —
#           reserved for last because layout is lost; the block is labeled so
#           the model knows. Scans have no text layer and are never swapped.
# Everything is best-effort per file: an unreadable/encrypted PDF is left
# untouched, and the caller still refuses the request if it cannot fit.

_SHRINK_MAX_DIM = 1600          # px — ~150-200 DPI for a letter page
_SHRINK_JPEG_QUALITY = 60
_PDF_TEXT_MIN_CHARS_PER_PAGE = 200   # below this the "text layer" is OCR noise
_PDF_TEXT_CAP = 150_000         # chars — one statement must not crowd the deal


def _recompress_pdf(raw: bytes) -> bytes | None:
    """Downscale and JPEG-re-encode every image inside a PDF.

    Returns the smaller PDF, or None when nothing changed or the result was not
    actually smaller (a CCITT fax-compressed scan can GROW as JPEG — keep the
    original then). Undecodable images are skipped, not fatal.
    """
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter(clone_from=PdfReader(io.BytesIO(raw)))
    changed = False
    for page in writer.pages:
        for img in page.images:
            try:
                pil = img.image
                if pil is None:
                    continue
                w, h = pil.size
                scale = _SHRINK_MAX_DIM / max(w, h)
                if scale < 1.0:
                    pil = pil.resize((max(1, round(w * scale)),
                                      max(1, round(h * scale))))
                if pil.mode not in ("RGB", "L"):
                    pil = pil.convert("RGB")
                img.replace(pil, quality=_SHRINK_JPEG_QUALITY)
                changed = True
            except Exception:  # noqa: BLE001 - leave this image as it is
                continue
    if not changed:
        return None
    out = io.BytesIO()
    writer.write(out)
    slim = out.getvalue()
    return slim if len(slim) < len(raw) else None


def _pdf_text(raw: bytes) -> str:
    """The PDF's text layer, or "" when it is too thin to stand in for the
    document (a scan). Capped so one long statement cannot crowd out the deal."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
        total += len(t)
        if total > _PDF_TEXT_CAP:
            parts.append(f"[... truncated at {_PDF_TEXT_CAP} characters ...]")
            break
    text = "\n".join(parts).strip()
    if len(text) < len(reader.pages) * _PDF_TEXT_MIN_CHARS_PER_PAGE:
        return ""
    return text[:_PDF_TEXT_CAP + 100]


def _wire_len(block: dict) -> int:
    if block.get("type") == "text":
        return len(block.get("text") or "")
    return len((block.get("source") or {}).get("data") or "")


def shrink_blocks_to_fit(docs: list[UploadedDoc], blocks: list[dict],
                         budget_bytes: int) -> list[str]:
    """Shrink the largest blocks IN PLACE until the request fits the budget.

    Returns one human-readable note per change (empty when nothing was done) —
    the caller logs them so a degraded document is visible, never silent.
    Never raises: a file that cannot be shrunk is simply left as uploaded.
    """
    notes: list[str] = []

    def total() -> int:
        return sum(_wire_len(b) for b in blocks)

    def mb(n: int) -> str:
        return f"{n / 1_000_000:.1f} MB"

    def by_size():
        pairs = [(d, b) for d, b in zip(docs, blocks)
                 if b.get("type") == "document"]
        return sorted(pairs, key=lambda db: _wire_len(db[1]), reverse=True)

    # Pass 1 — recompress the images inside the largest PDFs.
    for doc, block in by_size():
        if total() <= budget_bytes:
            break
        try:
            before = _wire_len(block)
            slim = _recompress_pdf(base64.b64decode(block["source"]["data"]))
            if slim is None:
                continue
            data = base64.b64encode(slim).decode("ascii")
            if len(data) >= before:
                continue
            block["source"]["data"] = data
            notes.append(f"{doc.filename or '(unnamed)'}: images recompressed, "
                         f"{mb(before)} -> {mb(len(data))} encoded")
        except Exception:  # noqa: BLE001 - an unreadable PDF stays as uploaded
            continue

    # Pass 2 — swap what still does not fit for its text layer (born-digital
    # statements collapse to a few hundred KB; scans return "" and stay put).
    for doc, block in by_size():
        if total() <= budget_bytes:
            break
        try:
            text = _pdf_text(base64.b64decode(block["source"]["data"]))
            if not text:
                continue
            before = _wire_len(block)
            name = doc.filename or "(unnamed file)"
            block.clear()
            block.update({
                "type": "text",
                "text": (f"=== {name} (text extracted from the PDF — the upload "
                         f"exceeded the API size limit, so the page images were "
                         f"not sent) ===\n{text}"),
            })
            notes.append(f"{name}: sent as its text layer, "
                         f"{mb(before)} -> {mb(_wire_len(block))}")
        except Exception:  # noqa: BLE001
            continue

    return notes


def build_document_blocks(docs: list[UploadedDoc]) -> list[dict]:
    """Build the content blocks for an extraction call, one per uploaded file.

    Raises RuntimeError naming the offending file when a document cannot be
    sent in any form — the caller surfaces that message to the user.
    """
    blocks: list[dict] = []
    for doc in docs:
        head = _head(doc)
        ext = _ext(doc)
        name = doc.filename or "(unnamed file)"

        if head.startswith(b"%PDF"):
            blocks.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf",
                           "data": doc.b64},
            })
            continue

        image_type = next(
            (mt for magic, mt in _IMAGE_MAGIC if head.startswith(magic)), None)
        if image_type is None and head[:4] == b"RIFF" and head[8:12] == b"WEBP":
            image_type = "image/webp"
        if image_type:
            blocks.append({
                "type": "image",
                "source": {"type": "base64", "media_type": image_type,
                           "data": doc.b64},
            })
            continue

        # Office formats are ZIP containers — the extension tells them apart.
        if head.startswith(b"PK\x03\x04"):
            if ext == ".docx":
                text, kind = _docx_text(_raw(doc)), "Word document"
            elif ext in {".xlsx", ".xlsm"}:
                text, kind = _xlsx_text(_raw(doc)), "Excel workbook"
            else:
                raise RuntimeError(
                    f"{name} is a zipped Office file this app cannot read "
                    f"({ext or 'no extension'}). Save it as PDF and upload again.")
            blocks.append({
                "type": "text",
                "text": f"=== {name} (text extracted from the {kind}) ===\n{text}",
            })
            continue

        if ext in _TEXT_EXTS or (doc.mime or "").startswith("text/"):
            blocks.append({
                "type": "text",
                "text": f"=== {name} ===\n{_plain_text(_raw(doc), ext)}",
            })
            continue

        if ext in {".doc", ".xls", ".ppt"}:
            raise RuntimeError(
                f"{name} is in the legacy Office format ({ext}), which cannot be "
                f"read directly. Open it and save as PDF or {ext}x, then upload again.")

        raise RuntimeError(
            f"{name} is not a file type this app can read"
            f"{f' ({ext})' if ext else ''}. Upload it as PDF, an image, "
            f"Word/Excel, or plain text.")

    return blocks
