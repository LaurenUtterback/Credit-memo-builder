"""Uploads must become content blocks the Anthropic API actually accepts.

The bug these lock in: every uploader passed the browser's reported MIME type
into a document block, and the API accepts only 'application/pdf' there. Files
off a network share often arrive as 'application/octet-stream' or with no MIME
at all, so a perfectly good PDF could fail the entire extraction with

    messages.0.content.3.document.source.base64.media_type:
        Input should be 'application/pdf'

The file's magic bytes decide, never the declared MIME.
"""

import base64
import io

import pytest

from app.models import UploadedDoc
from app.doc_blocks import build_document_blocks


def _doc(filename: str, data: bytes, mime: str = "") -> UploadedDoc:
    return UploadedDoc(filename=filename, mime=mime,
                       b64=base64.b64encode(data).decode())


MINIMAL_PDF = (b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\ntrailer<</Root 1 0 R>>\n"
               b"%%EOF\n")
PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmM"
    "IQAAAABJRU5ErkJggg==")


@pytest.mark.parametrize("mime", ["application/pdf", "application/octet-stream", ""])
def test_pdf_is_sent_as_a_document_whatever_the_browser_claims(mime):
    blocks = build_document_blocks([_doc("PFS.pdf", MINIMAL_PDF, mime)])
    assert blocks == [{
        "type": "document",
        "source": {"type": "base64", "media_type": "application/pdf",
                   "data": base64.b64encode(MINIMAL_PDF).decode()},
    }]


def test_image_is_sent_as_an_image_block_not_a_document():
    blocks = build_document_blocks([_doc("scan.png", PNG_1PX, "application/octet-stream")])
    assert blocks[0]["type"] == "image"
    assert blocks[0]["source"]["media_type"] == "image/png"


def test_word_document_is_converted_to_text():
    import docx
    d = docx.Document()
    d.add_paragraph("Term Sheet")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Loan Amount"
    table.rows[0].cells[1].text = "$2,025,000"
    buf = io.BytesIO()
    d.save(buf)

    blocks = build_document_blocks([_doc("terms.docx", buf.getvalue())])
    assert blocks[0]["type"] == "text"
    assert "Term Sheet" in blocks[0]["text"]
    assert "Loan Amount | $2,025,000" in blocks[0]["text"]
    assert "terms.docx" in blocks[0]["text"]      # Claude is told the source


def test_excel_workbook_is_converted_to_text():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Balloon"
    ws.append(["Gross Loan Amount", 2_025_000])
    ws.append([None, None])                        # blank rows are skipped
    ws.append(["Origination Fee", 60_750])
    buf = io.BytesIO()
    wb.save(buf)

    text = build_document_blocks([_doc("fees.xlsx", buf.getvalue())])[0]["text"]
    assert "sheet: Balloon" in text
    assert "Gross Loan Amount | 2025000" in text
    assert "Origination Fee | 60750" in text


def test_csv_and_plain_text_are_inlined():
    csv_text = build_document_blocks(
        [_doc("payoffs.csv", b"Lender,Balance\nNorthgate Bank,100000\n")])[0]["text"]
    assert "Lender | Balance" in csv_text
    assert "Northgate Bank | 100000" in csv_text

    txt = build_document_blocks([_doc("notes.txt", b"maturity 5/15/2031")])[0]["text"]
    assert "maturity 5/15/2031" in txt


def test_several_documents_keep_their_upload_order():
    blocks = build_document_blocks([
        _doc("a.pdf", MINIMAL_PDF),
        _doc("b.png", PNG_1PX),
        _doc("c.txt", b"hello"),
    ])
    assert [b["type"] for b in blocks] == ["document", "image", "text"]


def test_legacy_office_file_names_itself_in_the_error():
    # The old code turned this into a raw API 400 that named no file.
    with pytest.raises(RuntimeError) as exc:
        build_document_blocks([_doc("Old Term Sheet.doc", b"\xd0\xcf\x11\xe0junk")])
    assert "Old Term Sheet.doc" in str(exc.value)
    assert "PDF" in str(exc.value)


def test_unreadable_file_names_itself_in_the_error():
    with pytest.raises(RuntimeError) as exc:
        build_document_blocks([_doc("weird.dat", b"\x00\x01\x02\x03binary")])
    assert "weird.dat" in str(exc.value)
